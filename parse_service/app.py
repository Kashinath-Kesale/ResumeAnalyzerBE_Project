from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
import os
from io import BytesIO
from pdfminer.high_level import extract_text
from sklearn.feature_extraction.text import TfidfVectorizer
import tempfile
from docx import Document

ALLOWED = {'.pdf', '.doc', '.docx', '.txt'}
app = Flask(__name__)

def extract_text_from_file(path, ext):
    """Extract text from various file types"""
    if ext == '.pdf':
        return extract_text(path)
    elif ext == '.txt':
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.docx':
        # Extract text from DOCX using python-docx
        doc = Document(path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += '\n' + cell.text
        return text
    elif ext == '.doc':
        # .doc (old Word format) is not supported by python-docx
        # Would need additional library like antiword or textract
        raise ValueError("Legacy .doc format is not supported. Please convert to .docx or PDF")
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def extract_keywords(text, topn=20):
    # naive TF-IDF keyword extraction across single document:
    vec = TfidfVectorizer(stop_words='english', ngram_range=(1,2), max_features=2000)
    tfidf = vec.fit_transform([text])
    feature_names = vec.get_feature_names_out()
    # get top-n by tfidf score
    scores = tfidf.toarray()[0]
    idx = scores.argsort()[::-1][:topn]
    keywords = [feature_names[i] for i in idx if scores[i] > 0]
    return keywords

@app.route("/parse", methods=["POST"])
def parse():
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "no file"}), 400
    filename = secure_filename(f.filename)
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED:
        return jsonify({"error": "unsupported"}), 400
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
    f.save(tmp.name)
    try:
        text = extract_text_from_file(tmp.name, ext)
        if not text or not text.strip():
            return jsonify({"parsedText": "", "keywords": [], "error": "No text could be extracted from the file"})
        keywords = extract_keywords(text, topn=30)
        return jsonify({"parsedText": text, "keywords": keywords})
    except Exception as e:
        return jsonify({"parsedText": "", "keywords": [], "error": str(e)})
    finally:
        try:
            os.unlink(tmp.name)
        except:
            pass

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000)
